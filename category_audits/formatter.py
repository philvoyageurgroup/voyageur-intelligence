"""Markdown → branded DOCX formatter for Category Audit reports.

Parses structured markdown from Claude analysis and produces a
Voyageur-branded DOCX document with cover page, styled headings,
tables, bullet lists, and footer page numbers.
"""

from __future__ import annotations

import os
import re
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .models import CategoryAuditData

# ---------------------------------------------------------------------------
# Brand colours
# ---------------------------------------------------------------------------

NAVY = RGBColor(0x1F, 0x38, 0x64)       # #1F3864 — headings, table header
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)  # body text
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2) # table alt row
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FONT_BODY = "Calibri"
FONT_HEADING = "Calibri"

# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------


def _setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    return doc


def _add_logo_header(doc: Document):
    """Add Voyageur logo to header (top right). Skip if logo not found."""
    # Look for logo: first try 2 levels up (handoff/assets/), then 3 levels (src/ layout)
    base_2 = os.path.join(os.path.dirname(os.path.dirname(__file__)), "assets", "voyageur_logo.png")
    base_3 = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "assets", "voyageur_logo.png")
    logo_path = base_2 if os.path.exists(base_2) else base_3
    if not os.path.exists(logo_path):
        return
    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.add_picture(logo_path, width=Inches(1.5))


def _add_page_numbers(doc: Document):
    """Add centered page numbers to footer."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    _set_run_font(run, FONT_BODY, Pt(9), DARK_GRAY)

    # PAGE field
    fld_char_begin = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    run._element.append(fld_char_begin)
    instr = run._element.makeelement(qn("w:instrText"), {})
    instr.text = " PAGE "
    run._element.append(instr)
    fld_char_end = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run._element.append(fld_char_end)


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------


def _add_cover_page(doc: Document, data: CategoryAuditData):
    """Add title/cover page with category name, type, date."""
    # Spacer
    for _ in range(6):
        doc.add_paragraph("")

    # Title
    title = data.subcategory_name or data.category_name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    _set_run_font(run, FONT_HEADING, Pt(28), NAVY, bold=True)

    # Subtitle — report type
    type_labels = {
        "prospect": "Subcategory Intelligence Report",
        "brand": "Brand Health Report",
        "buyer": "Buyer Intelligence Report",
    }
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(type_labels.get(data.report_type, "Category Report"))
    _set_run_font(run, FONT_HEADING, Pt(18), DARK_GRAY)

    # Target brand or retailer
    if data.target_brand:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Target Brand: {data.target_brand}")
        _set_run_font(run, FONT_BODY, Pt(14), DARK_GRAY)

    if data.retailer:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Prepared for: {data.retailer}")
        _set_run_font(run, FONT_BODY, Pt(14), DARK_GRAY)

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(data.data_pulled_at.strftime("%B %Y"))
    _set_run_font(run, FONT_BODY, Pt(12), DARK_GRAY)

    # Spacer
    doc.add_paragraph("")

    # Prepared by
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared by Voyageur Group")
    _set_run_font(run, FONT_BODY, Pt(11), DARK_GRAY, italic=True)

    # Page break
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Markdown parser → DOCX
# ---------------------------------------------------------------------------


def _strip_leading_title(markdown: str) -> str:
    """Remove the first heading block if it looks like a duplicate title.

    Claude often generates a '# Category Intelligence Report: ...' line
    at the very top of the markdown.  Since we already have a cover page
    with identical information, strip it to avoid duplication.

    Removes the first ``# ...`` heading (and any immediately following
    blank lines or a single subtitle-like line) before the first ``##``.
    """
    lines = markdown.split("\n")
    # Find first non-blank line
    start = 0
    while start < len(lines) and not lines[start].strip():
        start += 1

    if start >= len(lines):
        return markdown

    first = lines[start].strip()
    # Only strip if it's a top-level heading (single #)
    if first.startswith("# ") and not first.startswith("## "):
        # Drop that line
        start += 1
        # Drop any blank lines or a single non-heading subtitle right after
        while start < len(lines):
            s = lines[start].strip()
            if not s:
                start += 1
                continue
            if s.startswith("#"):
                break  # hit next real section — stop stripping
            # It might be a subtitle like "SpaceAid Prospect Analysis" — skip it
            start += 1
            break  # only skip one subtitle line
        # Skip trailing blanks after subtitle
        while start < len(lines) and not lines[start].strip():
            start += 1
        return "\n".join(lines[start:])

    return markdown


def _parse_and_render(doc: Document, markdown: str):
    """Parse markdown and render into DOCX elements."""
    markdown = _strip_leading_title(markdown)
    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule → thin line / page break hint
        if stripped in ("---", "***", "___"):
            _add_hr(doc)
            i += 1
            continue

        # Heading: ## or ###
        if stripped.startswith("###"):
            text = stripped.lstrip("#").strip()
            _add_heading(doc, text, level=2)
            i += 1
            continue
        if stripped.startswith("##"):
            text = stripped.lstrip("#").strip()
            _add_heading(doc, text, level=1)
            i += 1
            continue
        if stripped.startswith("#"):
            text = stripped.lstrip("#").strip()
            _add_heading(doc, text, level=0)
            i += 1
            continue

        # Table: | ... |
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _add_table(doc, table_lines)
            continue

        # Bullet list: - or * or • (Claude sometimes uses Unicode bullets)
        if stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("\u2022 "):
            bullet_lines = []
            while i < len(lines):
                s = lines[i].strip()
                if s.startswith("- ") or s.startswith("* ") or s.startswith("\u2022 "):
                    bullet_lines.append(s[2:])
                    i += 1
                elif s.startswith("  ") and bullet_lines:
                    # continuation of previous bullet
                    bullet_lines[-1] += " " + s.strip()
                    i += 1
                else:
                    break
            for b in bullet_lines:
                _add_bullet(doc, b)
            continue

        # Numbered list: 1. 2. etc.
        num_match = re.match(r"^(\d+)\.\s+", stripped)
        if num_match:
            num_lines = []
            while i < len(lines):
                s = lines[i].strip()
                nm = re.match(r"^(\d+)\.\s+", s)
                if nm:
                    num_lines.append(s[nm.end():])
                    i += 1
                elif s.startswith("  ") and num_lines:
                    num_lines[-1] += " " + s.strip()
                    i += 1
                else:
                    break
            for idx, nl in enumerate(num_lines, 1):
                _add_numbered(doc, nl, idx)
            continue

        # Regular paragraph
        _add_paragraph(doc, stripped)
        i += 1


# ---------------------------------------------------------------------------
# Element builders
# ---------------------------------------------------------------------------


def _set_run_font(
    run,
    font_name: str = FONT_BODY,
    size: Pt = Pt(11),
    color: RGBColor = DARK_GRAY,
    bold: bool = False,
    italic: bool = False,
):
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def _add_heading(doc: Document, text: str, level: int = 1):
    """Add a styled heading using real Word heading styles.

    Level 0 = Heading 1 (title-ish, ##-less top-level)
    Level 1 = Heading 1 (## in markdown)
    Level 2 = Heading 2 (### in markdown)

    Uses doc.add_heading() so Word's navigation pane and
    style dropdown show proper Heading 1 / Heading 2.
    """
    # Map our levels to Word heading levels (1-based)
    word_level = 1 if level <= 1 else 2
    sizes = {1: Pt(16), 2: Pt(13)}

    heading = doc.add_heading(level=word_level)
    heading.space_before = Pt(12) if level > 0 else Pt(6)
    heading.space_after = Pt(4)
    # Clear default run that add_heading may create, render with our styling
    _render_inline(heading, text, font_name=FONT_HEADING,
                   size=sizes.get(word_level, Pt(14)),
                   color=NAVY, default_bold=True)


def _add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    _render_inline(p, text)


def _add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.space_after = Pt(2)
    _render_inline(p, text)


def _add_numbered(doc: Document, text: str, num: int):
    p = doc.add_paragraph(style="List Number")
    p.space_after = Pt(2)
    _render_inline(p, text)


def _add_hr(doc: Document):
    """Add a subtle horizontal rule (thin paragraph border)."""
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "1",
            qn("w:color"): "CCCCCC",
        },
    )
    pBdr.append(bottom)
    pPr.append(pBdr)


def _render_inline(
    p,
    text: str,
    font_name: str = FONT_BODY,
    size: Pt = Pt(11),
    color: RGBColor = DARK_GRAY,
    default_bold: bool = False,
):
    """Parse inline markdown (**bold**, *italic*) and render as runs."""
    # Split on **bold** and *italic* markers
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            _set_run_font(run, font_name, size, color, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            _set_run_font(run, font_name, size, color, italic=True)
        else:
            run = p.add_run(part)
            _set_run_font(run, font_name, size, color, bold=default_bold)


def _add_table(doc: Document, lines: List[str]):
    """Parse markdown table and render as DOCX table."""
    # Parse rows
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        # Skip separator row (---|---|---)
        if all(re.match(r"^-+:?$|^:?-+:?$", c) for c in cells):
            continue
        rows.append(cells)

    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= n_cols:
                break
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.space_after = Pt(0)
            p.space_before = Pt(0)

            is_header = r_idx == 0
            _render_inline(
                p,
                cell_text,
                size=Pt(9),
                color=WHITE if is_header else DARK_GRAY,
                default_bold=is_header,
            )

            # Header row: navy background
            if is_header:
                _set_cell_bg(cell, "1F3864")
            elif r_idx % 2 == 0:
                _set_cell_bg(cell, "F2F2F2")


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.makeelement(
        qn("w:shd"),
        {
            qn("w:fill"): hex_color,
            qn("w:val"): "clear",
        },
    )
    tcPr.append(shading)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def generate_docx(
    markdown: str,
    data: CategoryAuditData,
    output_dir: str = "output/category_audits/",
) -> str:
    """Generate a branded DOCX from markdown analysis + metadata.

    Returns the output file path.
    """
    doc = _setup_document()
    _add_logo_header(doc)
    _add_page_numbers(doc)
    _add_cover_page(doc, data)
    _parse_and_render(doc, markdown)

    os.makedirs(output_dir, exist_ok=True)
    ts = data.data_pulled_at.strftime("%Y%m%d_%H%M")
    safe_name = (data.subcategory_name or data.category_name).replace(" ", "_").lower()
    filename = f"{safe_name}_{data.report_type}_{ts}.docx"
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    return output_path
